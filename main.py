import asyncio
import re
import os
from datetime import datetime
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
import openai
from io import BytesIO
import logging
from example import collect_product_data
from dotenv import load_dotenv
from aiogram.utils.keyboard import InlineKeyboardMarkup, InlineKeyboardButton
from fpdf import FPDF
from aiogram.types import FSInputFile
from docx import Document
from docx2pdf import convert
import subprocess

load_dotenv()

# Настройки API
API_TOKEN = os.getenv('API_TOKEN')

PROXY_API_KEY = os.getenv("OPENAI_API_KEY")

# Настройка OpenAI с ProxyAPI
openai.api_key = PROXY_API_KEY
openai.api_base = "https://api.proxyapi.ru/openai/v1"

bot = Bot(token=API_TOKEN)
dp = Dispatcher()

logging.basicConfig(level=logging.INFO)
global_advantages_text = ""
import requests


async def extract_and_format_text_from_image(img_urls_list):
    # Фильтруем пустые строки и ограничиваем количество до 18 изображений
    img_urls = [url for url in img_urls_list if url][:18]

    # Загружаем текстовый пример из файла
    with open("image_example", "r", encoding="utf-8") as text_file:
        example = text_file.read()

    # Проверяем наличие изображений
    if not img_urls:
        return "No valid URLs found."

    # Формируем список контента для сообщений с URL изображениями
    messages_content = [
        {"type": "text",
         "text": "Распознай весь текст на изображениях по следующим URL и подробно опиши по каждому продукту."}
    ]
    messages_content.extend(
        [{"type": "image_url", "image_url": {"url": url}} for url in img_urls]
    )

    # Формируем и выполняем запрос к API
    response = await openai.ChatCompletion.acreate(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": f"Распознай весь текст на изображениях, без перефразирования, а также с распределением по слайдам, убирая все знаки ** и ###. Примеры: {example}"
            },
            {
                "role": "user",
                "content": messages_content,
            }
        ],
        max_tokens=800
    )

    # Возвращаем текст ответа, если он есть
    return response.choices[0].message.content.strip()    # Фильтруем пустые строки и ограничиваем количество до 18




# Функция для генерации преимуществ товара с использованием OpenAI
async def generate_product_advantages(description, formatted_texts, product_name):
    global num_slides
    with open("text_example", "r", encoding="utf-8") as text_file:
        example = text_file.read()
    prompt = (

        f"Сгенерируй техническое задание с преимуществами товара на основе описания и текста с изображений:\n\n"
        f"Нужно обязательно убирать знаки ###, ** \n\n"
        f"Название продукта: Выведи из всех описаний общее название"
        f"Описание: {description}\n\n"
        f"Текст с изображений: {formatted_texts}\n\n"
        f"Выведи сначала название товара, потом преимущества товара(от 20 до 30 преимуществ и более) и техническое задание по слайдам (количество слайдов {num_slides}) в структурированном виде!"
        f"Пример конечного результата: {example}, количество слайдов равно {num_slides}"
    )

    # Использование ChatCompletion для gpt-4 модели
    response = await openai.ChatCompletion.acreate(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": prompt},
                  {"role": "user", "content": prompt}],
        max_tokens=1500
    )
    return response['choices'][0]['message']['content'].strip()


# Обработчик команды /start
num_slides = 0  # Глобальная переменная для хранения количества слайдов


@dp.message(Command("start"))
async def send_welcome(message: types.Message):
    markup = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Выберите количество слайдов", callback_data="distribute_slides")]
    ])
    await message.answer("Привет! Для начала выберите количество слайдов:", reply_markup=markup)


@dp.callback_query(lambda c: c.data == "distribute_slides")
async def distribute_slides(callback_query: types.CallbackQuery):
    buttons = [[InlineKeyboardButton(text=str(i), callback_data=f"slides_{i}") for i in range(row, row + 5)] for row in
               range(1, 11, 5)]
    keyboard = InlineKeyboardMarkup(inline_keyboard=buttons)
    await callback_query.message.answer("Выберите количество слайдов:", reply_markup=keyboard)


@dp.callback_query(lambda c: c.data.startswith("slides_"))
async def handle_slide_selection(callback_query: types.CallbackQuery):
    global num_slides
    num_slides = int(callback_query.data.split("_")[1])
    await callback_query.message.answer(
        f"Вы выбрали {num_slides} слайдов. Теперь введите ссылку(ссылки) на товары Ozon или Wildberries.")


def split_message(message, limit=4096):
    for i in range(0, len(message), limit):
        yield message[i:i + limit]


@dp.message(lambda message: re.search(r'^https://www\.(wildberries|ozon)\.ru', message.text))
async def handle_multiple_links(message: types.Message):
    global num_slides
    links = message.text.split()
    if len(links) > 5:
        await message.answer("Пожалуйста, отправьте не более 5 ссылок.")
        return

    await message.answer("Начали сбор данных по товарам. Подождите 2 минуты.")

    # Collect all descriptions and images across links
    all_descriptions = []
    all_images = []
    product_name = "Товар"  # Default name if none is found

    for url in links:
        product_data = collect_product_data(url)
        all_descriptions.append(product_data.get('description', "Описание не найдено"))
        all_images.extend(product_data.get('images', []))  # Add all images from each product
        if not product_name and "name" in product_data:
            product_name = product_data.get("name", "Товар")  # Set name from the first product if available

    # Join all descriptions into one
    combined_description = " ".join(all_descriptions)

    # Ensure unique images only and limit to 18 images for processing
    unique_images = list(dict.fromkeys(all_images))[:18]
    formatted_texts = await extract_and_format_text_from_image(unique_images)

    # Generate advantages text based on combined data
    advantages_text = await generate_product_advantages(combined_description, formatted_texts, product_name)
    await message.answer(formatted_texts)
    for part in split_message(advantages_text):
        await message.answer(part)
    # await message.answer(f"Текст с картинок:\n{formatted_texts}\n\nТекст с описания:\n{combined_description}")

    # Store the global advantages text and generate the final document
    global global_advantages_text
    global_advantages_text = advantages_text
    await generate_document(message, product_name)

index=0
# Функция для создания PDF документа
async def generate_document(message, product_name):
    global global_advantages_text, num_slides, index

    # Clean up `product_name` by removing unsuitable characters
    product_name = re.sub(r'[\\/*?:"<>|\n\t]', "", product_name).strip()

    # Generate a unique filename with incrementing numbers
    base_filename = product_name or "Товар"
    index += 1

    unique_name = f"{base_filename}_{index}"

    # File paths for DOCX and PDF
    temp_docx_file = f"{unique_name}.docx"
    temp_pdf_file = f"{unique_name}.pdf"

    # Create a new Word document
    doc = Document()
    doc.add_heading("Техническое задание для слайдов", level=1)
    doc.add_paragraph(f"Количество слайдов: {num_slides}")
    doc.add_paragraph(global_advantages_text)

    # Save the document to a temporary DOCX file
    doc.save(temp_docx_file)

    # Convert the DOCX file to a PDF
    convert(temp_docx_file, temp_pdf_file)

    # Send the generated PDF file to the user
    pdf_file = FSInputFile(temp_pdf_file)
    await message.answer_document(pdf_file)

    # Remove the temporary DOCX and PDF files
    os.remove(temp_docx_file)
    os.remove(temp_pdf_file)

    # Restore the menu for the user
    markup = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Выберите количество слайдов", callback_data="distribute_slides")]
    ])
    await message.answer("Выберите количество слайдов:", reply_markup=markup)


@dp.message(F.text)
async def process_message(message: types.Message):
    await message.answer(
        "Неверный формат сообщения. Вводите только ссылки."
    )
    return


async def main():
    await dp.start_polling(bot)


if __name__ == '__main__':
    asyncio.run(main())
